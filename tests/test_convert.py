"""Tests for mdup: rendering, fallbacks, and end-to-end conversion."""

import zipfile
from pathlib import Path

import pytest

from mdup import convert, render
from mdup.cli import main
from mdup.core import MarkdownConverter

FIXTURES = Path(__file__).parent / "fixtures"
SAMPLE = FIXTURES / "sample.md"


# ----------------------------------------------------------------- rendering

def test_render_covers_markdown_features(tmp_path):
    html = render.render_body(SAMPLE.read_text(), work_dir=tmp_path, toc=True)
    assert "<table>" in html                      # GFM table
    assert "<s>" in html                           # strikethrough
    assert "task-list-item" in html                # task list
    assert "footnote" in html.lower()              # footnotes
    assert 'id="lists"' in html or "id=\"lists\"" in html  # heading anchor
    assert 'class="toc"' in html                   # table of contents
    assert "title:" not in html                    # YAML front-matter stripped


def test_toc_marker_replaced(tmp_path):
    md = "# A\n\n[TOC]\n\n## Section One\n\n## Section Two\n"
    html = render.render_body(md, work_dir=tmp_path, toc=False)
    assert 'class="toc"' in html
    assert "[TOC]" not in html


# ----------------------------------------------------------------- fallbacks

def test_mermaid_fallback_when_no_renderer(tmp_path, monkeypatch):
    monkeypatch.setattr("mdup.mermaid._MMDC_PATH", False)
    monkeypatch.setattr("mdup.mermaid._WARNED", False)
    md = "```mermaid\nflowchart LR\n A-->B\n```\n"
    html = render.render_body(md, work_dir=tmp_path)
    assert 'class="mermaid"' in html               # styled code block, not an image
    assert "<img" not in html


def _png_header(width: int, height: int) -> bytes:
    """A minimal PNG signature + IHDR — enough for _png_size_pt to read its size."""
    import struct
    return b"\x89PNG\r\n\x1a\n" + b"\x00\x00\x00\x0dIHDR" + struct.pack(">II", width, height)


def test_mermaid_wide_image_capped_to_page_width(tmp_path):
    from mdup import mermaid

    wide = tmp_path / "wide.png"
    wide.write_bytes(_png_header(2000, 500))   # far wider than the page
    w_pt, h_pt = mermaid._png_size_pt(wide)
    assert abs(w_pt - mermaid._MAX_WIDTH_PT) < 0.05  # clamped to page width
    assert abs(h_pt - w_pt * 0.25) < 0.05            # aspect ratio (500/2000) preserved


def test_mermaid_small_image_not_upscaled(tmp_path):
    from mdup import mermaid

    small = tmp_path / "small.png"
    small.write_bytes(_png_header(96, 48))     # 96px @96dpi -> 72pt, well under the page
    w_pt, h_pt = mermaid._png_size_pt(small)
    assert (round(w_pt, 1), round(h_pt, 1)) == (72.0, 36.0)


def test_math_fallback_when_no_matplotlib(tmp_path, monkeypatch):
    monkeypatch.setattr("mdup.mathrender._MATPLOTLIB", False)
    monkeypatch.setattr("mdup.mathrender._WARNED", False)
    monkeypatch.setattr("mdup.mathrender._cache", {})
    html = render.render_body("Energy $E=mc^2$ here.", work_dir=tmp_path)
    assert "math-inline-fallback" in html          # literal text, not an image
    assert "E=mc^2" in html


# --------------------------------------------------------------- conversion

def test_convert_single_both_formats(tmp_path):
    out = tmp_path / "out"
    written = convert([SAMPLE], formats=["pdf", "docx"], output=out)
    assert len(written) == 2
    for p in written:
        assert p.exists() and p.stat().st_size > 0
    assert (out / "sample.pdf").exists()
    assert (out / "sample.docx").exists()


def test_convert_separate_multiple(tmp_path):
    a = tmp_path / "a.md"; a.write_text("# Alpha\n\nText A.")
    b = tmp_path / "b.md"; b.write_text("# Beta\n\nText B.")
    out = tmp_path / "sep"
    written = convert([a, b], formats=["docx"], output=out)
    assert {p.name for p in written} == {"a.docx", "b.docx"}


def test_merge_has_single_consolidated_toc(tmp_path):
    """Merging several files with --toc yields one TOC at the top, not one per file."""
    from docx import Document

    a = tmp_path / "a.md"; a.write_text("# Alpha\n\n## Intro\n\n## Methods\n")
    b = tmp_path / "b.md"; b.write_text("# Beta\n\n## Intro\n\n## Results\n")
    written = convert([a, b], formats=["docx"], output=tmp_path / "m.docx",
                      merge=True, toc=True)
    doc = Document(str(written[0]))
    titles = [p.text.strip() for p in doc.paragraphs]
    assert titles.count("Contents") == 1


def test_merge_toc_namespaces_duplicate_anchors(tmp_path):
    """A consolidated TOC links unambiguously when files share heading titles."""
    a = tmp_path / "a.md"; a.write_text("# Alpha\n\n## Intro\n")
    b = tmp_path / "b.md"; b.write_text("# Beta\n\n## Intro\n")
    conv = MarkdownConverter(["pdf"], merge=True, toc=True)
    rendered = [(p, *conv._render_body(p, tmp_path)) for p in (a, b)]
    parts, headings = [], []
    for i, (_, body, heads) in enumerate(rendered):
        body, heads = render.prefix_anchors(body, heads, f"f{i}-")
        headings.extend(heads)
    toc = render.build_toc(headings)
    assert 'href="#f0-intro"' in toc and 'href="#f1-intro"' in toc


def test_convert_merge_single_file(tmp_path):
    a = tmp_path / "a.md"; a.write_text("# Alpha")
    b = tmp_path / "b.md"; b.write_text("# Beta")
    out = tmp_path / "book.pdf"
    written = convert([a, b], formats=["pdf"], output=out, merge=True)
    assert len(written) == 1
    assert written[0] == out and out.exists()


def test_docx_oversized_image_scaled_to_fit(tmp_path):
    """An image wider than the page text area is scaled down, aspect ratio kept."""
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    from mdup.backends import _fit_images_to_page

    img = tmp_path / "wide.png"
    Image.new("RGB", (2000, 300), "white").save(img)
    doc = Document()
    doc.add_picture(str(img), width=Inches(12))  # force it past the ~6in text column
    _fit_images_to_page(doc)

    sec = doc.sections[0]
    max_w = sec.page_width - sec.left_margin - sec.right_margin
    shape = doc.inline_shapes[0]
    assert shape.width <= max_w
    assert abs(shape.width / shape.height - 2000 / 300) < 0.02  # aspect preserved


def test_docx_embeds_images(tmp_path):
    written = convert([SAMPLE], formats=["docx"], output=tmp_path)
    with zipfile.ZipFile(written[0]) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    assert media, "expected at least the local logo image to be embedded"


# ----------------------------------------------------------------- validation

def test_invalid_format_rejected():
    with pytest.raises(ValueError):
        MarkdownConverter(["txt"])


def test_missing_input_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        convert([tmp_path / "nope.md"], formats=["pdf"], output=tmp_path)


def test_non_utf8_input_reports_filename(tmp_path):
    """A binary/non-UTF-8 file (e.g. a glob that caught a .docx) names the offender."""
    bad = tmp_path / "binary.docx"
    bad.write_bytes(b"PK\x03\x04\x14\x00\x06\x00\x08\x00\xa4\x00binary")
    with pytest.raises(ValueError, match=r"binary\.docx"):
        convert([bad], formats=["docx"], output=tmp_path)


# ------------------------------------------------------------------------ CLI

def test_cli_smoke(tmp_path, capsys):
    rc = main([str(SAMPLE), "-f", "docx", "-o", str(tmp_path)])
    assert rc == 0
    assert (tmp_path / "sample.docx").exists()


def test_cli_bad_input_returns_error(tmp_path, capsys):
    rc = main([str(tmp_path / "missing.md"), "-o", str(tmp_path)])
    assert rc == 2
